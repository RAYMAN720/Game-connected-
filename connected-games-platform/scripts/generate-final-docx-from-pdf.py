from pathlib import Path
import shutil
import subprocess

from docx import Document
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
PDF_SOURCE = ROOT / "docs" / "PlayConnect Relazione Progetto.pdf"
FINAL_DIR = ROOT / "docs" / "final"
FINAL_PDF = FINAL_DIR / "PlayConnect_Relazione_Finale.pdf"
FINAL_DOCX = FINAL_DIR / "PlayConnect_Relazione_Finale.docx"
TMP_DIR = ROOT / "tmp" / "pdfs" / "relazione-final-pages"
POPPLER = Path(
    r"C:\Users\nguet\.cache\codex-runtimes\codex-primary-runtime\dependencies\native\poppler\Library\bin\pdftoppm.exe"
)


def clear_pngs():
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    for old in TMP_DIR.glob("page-*.png"):
        old.unlink()


def render_pdf_pages():
    clear_pngs()
    subprocess.run(
        [str(POPPLER), "-r", "180", "-png", str(PDF_SOURCE), str(TMP_DIR / "page")],
        check=True,
    )
    return sorted(TMP_DIR.glob("page-*.png"))


def build_docx(page_images):
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.2)
    section.bottom_margin = Cm(0.2)
    section.left_margin = Cm(0.2)
    section.right_margin = Cm(0.2)

    for index, image in enumerate(page_images):
        if index:
            document.add_page_break()
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        run = paragraph.add_run()
        run.add_picture(str(image), width=Cm(20.6))

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    document.save(FINAL_DOCX)


def main():
    if not PDF_SOURCE.exists():
        raise FileNotFoundError(PDF_SOURCE)
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(PDF_SOURCE, FINAL_PDF)
    images = render_pdf_pages()
    if not images:
        raise RuntimeError("Nessuna pagina renderizzata dal PDF.")
    build_docx(images)
    print(FINAL_PDF)
    print(FINAL_DOCX)


if __name__ == "__main__":
    main()

from pathlib import Path
import re
import html

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Preformatted


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "00-relazione-finale.md"
PDF_MAIN = ROOT / "docs" / "PlayConnect Relazione Progetto.pdf"
FINAL_DIR = ROOT / "docs" / "final"
FINAL_PDF = FINAL_DIR / "PlayConnect_Relazione_Finale.pdf"
FINAL_DOCX = FINAL_DIR / "PlayConnect_Relazione_Finale.docx"


def clean_inline(text):
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escaped)
    return escaped


def markdown_blocks(text):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            yield ("code", "\n".join(code))
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table.append(lines[i].rstrip())
                i += 1
            yield ("table", table)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            yield (f"h{min(level, 3)}", line[level:].strip())
            i += 1
            continue
        if line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            yield ("ul", items)
            continue
        if re.match(r"^\d+\. ", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                items.append(re.sub(r"^\d+\. ", "", lines[i]).strip())
                i += 1
            yield ("ol", items)
            continue
        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "- ", "```", "|")) and not re.match(r"^\d+\. ", lines[i]):
            para.append(lines[i].strip())
            i += 1
        yield ("p", " ".join(para))


def build_pdf(blocks):
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleBlue", parent=styles["Title"], textColor=colors.HexColor("#1f4e79"), fontSize=24, leading=30, spaceAfter=16))
    styles.add(ParagraphStyle(name="H1Blue", parent=styles["Heading1"], textColor=colors.HexColor("#1f4e79"), fontSize=16, leading=20, spaceBefore=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2Blue", parent=styles["Heading2"], textColor=colors.HexColor("#2e75b6"), fontSize=13, leading=16, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="BodyJustified", parent=styles["BodyText"], fontSize=9.5, leading=12.5, spaceAfter=6))
    styles.add(ParagraphStyle(name="BulletSmall", parent=styles["BodyText"], fontSize=9.2, leading=12, leftIndent=14, firstLineIndent=-8, spaceAfter=3))
    def make_story():
        story = []
        for kind, value in blocks:
            if kind == "h1":
                story.append(Paragraph(clean_inline(value), styles["TitleBlue"]))
            elif kind == "h2":
                story.append(Paragraph(clean_inline(value), styles["H1Blue"]))
            elif kind == "h3":
                story.append(Paragraph(clean_inline(value), styles["H2Blue"]))
            elif kind == "p":
                story.append(Paragraph(clean_inline(value), styles["BodyJustified"]))
            elif kind == "ul":
                for item in value:
                    story.append(Paragraph(f"- {clean_inline(item)}", styles["BulletSmall"]))
            elif kind == "ol":
                for idx, item in enumerate(value, 1):
                    story.append(Paragraph(f"{idx}. {clean_inline(item)}", styles["BulletSmall"]))
            elif kind == "code":
                story.append(Preformatted(value, styles["Code"]))
                story.append(Spacer(1, 0.2 * cm))
            elif kind == "table":
                rows = []
                for row in value:
                    if re.match(r"^\|\s*-", row):
                        continue
                    cells = [cell.strip() for cell in row.strip("|").split("|")]
                    rows.append([Paragraph(clean_inline(cell), styles["BodyJustified"]) for cell in cells])
                if rows:
                    table = Table(rows, repeatRows=1, hAlign="LEFT")
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#17365d")),
                        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#b7c9d6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.25 * cm))
        return story

    for target in [PDF_MAIN, FINAL_PDF]:
        doc = SimpleDocTemplate(str(target), pagesize=A4, rightMargin=1.6 * cm, leftMargin=1.6 * cm, topMargin=1.6 * cm, bottomMargin=1.5 * cm)
        doc.build(make_story())


def build_docx(blocks):
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for kind, value in blocks:
        if kind == "h1":
            document.add_heading(value, level=0)
        elif kind == "h2":
            document.add_heading(value, level=1)
        elif kind == "h3":
            document.add_heading(value, level=2)
        elif kind == "p":
            document.add_paragraph(re.sub(r"[`*]", "", value))
        elif kind == "ul":
            for item in value:
                document.add_paragraph(re.sub(r"[`*]", "", item), style="List Bullet")
        elif kind == "ol":
            for item in value:
                document.add_paragraph(re.sub(r"[`*]", "", item), style="List Number")
        elif kind == "code":
            p = document.add_paragraph()
            run = p.add_run(value)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "table":
            rows = []
            for row in value:
                if re.match(r"^\|\s*-", row):
                    continue
                rows.append([cell.strip().replace("`", "") for cell in row.strip("|").split("|")])
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table.cell(r, c).text = cell
    document.save(FINAL_DOCX)


def main():
    text = SOURCE.read_text(encoding="utf-8")
    blocks = list(markdown_blocks(text))
    build_pdf(blocks)
    build_docx(blocks)
    print(PDF_MAIN)
    print(FINAL_PDF)
    print(FINAL_DOCX)


if __name__ == "__main__":
    main()

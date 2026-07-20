from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / 'docs' / 'figures'
OUT = ROOT / 'docs' / 'final' / 'PlayConnect_Diagrammi.docx'

items = [
    ('01-casi-uso.png', 'Diagramma dei casi d’uso'),
    ('02-architettura.png', 'Architettura generale'),
    ('03-dominio.png', 'Diagramma delle classi del dominio'),
    ('04-package.png', 'Diagramma dei package'),
    ('05-classi-implementazione.png', 'Classi e moduli di implementazione'),
    ('06-sequenza-online.png', 'Diagramma di sequenza: partita online'),
    ('07-sequenza-offline.png', 'Diagramma di sequenza: coda offline e sincronizzazione'),
    ('08-deployment.png', 'Diagramma di deployment Docker Compose'),
]

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.top_margin = Cm(1.2)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(1.3)
sec.right_margin = Cm(1.3)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
styles['Title'].font.name = 'Arial'
styles['Title'].font.size = Pt(26)
styles['Title'].font.color.rgb = RGBColor(31, 78, 121)
styles['Heading 1'].font.name = 'Arial'
styles['Heading 1'].font.size = Pt(18)
styles['Heading 1'].font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('PlayConnect')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Diagrammi UML e architetturali')
r.bold = True
r.font.size = Pt(17)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Progetto di Laboratorio PISSIR – A.A. 2025/2026')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'N.'
table.rows[0].cells[1].text = 'Diagramma'
for i, (_, title) in enumerate(items, 1):
    cells = table.add_row().cells
    cells[0].text = str(i)
    cells[1].text = title

for index, (filename, title) in enumerate(items, 1):
    doc.add_page_break()
    h = doc.add_paragraph(style='Heading 1')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run(f'{index}. {title}')
    path = FIG / filename
    with Image.open(path) as im:
        width_px, height_px = im.size
    max_w = Cm(25.8)
    max_h = Cm(14.7)
    ratio = width_px / height_px
    width = max_w
    height = width / ratio
    if height > max_h:
        height = max_h
        width = height * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width, height=height)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f'Figura {index} – {title}')
    run.italic = True
    run.font.size = Pt(9)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

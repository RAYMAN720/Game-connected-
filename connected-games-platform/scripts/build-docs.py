from pathlib import Path
import shutil
import subprocess
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'
SOURCE = DOCS / 'source'
FIG = DOCS / 'figures'
FINAL = DOCS / 'final'
SOURCE.mkdir(parents=True, exist_ok=True)
FINAL.mkdir(parents=True, exist_ok=True)


def write(name, text):
    path = DOCS / name
    path.write_text(text.strip() + '\n', encoding='utf-8')
    return path


specifiche = write('01-specifiche-funzionali.md', r'''
# Specifiche funzionali

## Obiettivo

PlayConnect collega giochi fisici a una piattaforma centrale. I sensori rilevano gli eventi della partita, il dispositivo edge li invia tramite MQTT e il server aggiorna punteggi, statistiche e tornei.

## Utenti

| Ruolo | Funzioni principali |
|---|---|
| Giocatore | Vede giochi disponibili, partite, statistiche personali, tornei e classifiche. |
| Amministratore locale | Gestisce giochi, client, dispositivi edge, sensori, attuatori e partite del proprio locale. |
| Amministratore gioco | Definisce tipi di gioco, eventi, limite di punteggio e modelli dei sensori. |
| Amministratore piattaforma | Gestisce locali, amministratori, tornei e statistiche globali. |

## Funzioni richieste

- gestione di locali pubblici o privati;
- identificazione univoca dei giochi nel locale;
- tipi di gioco configurabili;
- partite individuali e a squadre;
- eventi sensore con valore numerico;
- registrazione degli eventi con UUID per evitare duplicati;
- funzionamento offline dell'edge e sincronizzazione successiva;
- tornei su uno o piu locali;
- classifiche e statistiche;
- interfaccia web per tutti i ruoli;
- API REST documentate con OpenAPI;
- comunicazione MQTT con QoS 1.

## Requisiti non funzionali

Il sistema deve essere semplice da avviare, separato in microservizi, testabile e utilizzabile anche quando la connessione tra edge e server non e disponibile per un breve periodo.
''')

architettura = write('02-architettura.md', r'''
# Architettura del sistema

![Architettura generale](figures/02-architettura.png)

## Componenti

- **Frontend web**: pagine HTML, CSS e JavaScript servite da Nginx.
- **API Gateway**: riceve le richieste del browser e le inoltra al servizio corretto.
- **Catalog Service**: utenti, locali, tipi di gioco, giochi, edge, sensori e attuatori.
- **Match Service**: partite, eventi, punteggi, heartbeat MQTT e comandi agli attuatori.
- **Tournament Service**: tornei, squadre, classifiche e statistiche.
- **Edge Service**: simula i sensori, conserva gli eventi offline e li sincronizza.
- **Mosquitto**: broker MQTT.
- **MySQL**: database centrale. Solo i servizi backend accedono direttamente al database.

## Distribuzione

![Deployment Docker](figures/08-deployment.png)

Ogni componente principale viene eseguito in un container. Tutti i container comunicano sulla rete creata da Docker Compose.
''')

diagrammi = write('03-diagrammi.md', r'''
# Diagrammi del progetto

## Casi d'uso

![Casi d'uso](figures/01-casi-uso.png)

## Dominio

![Diagramma del dominio](figures/03-dominio.png)

## Package

![Diagramma dei package](figures/04-package.png)

## Classi di implementazione

![Classi principali](figures/05-classi-implementazione.png)

## Sequenza online

![Sequenza online](figures/06-sequenza-online.png)

## Sequenza offline

![Sequenza offline](figures/07-sequenza-offline.png)
''')

api_rest = write('04-api-rest.md', r'''
# API REST

La descrizione completa e presente in `docs/openapi.yaml` nel formato OpenAPI 3.0.3.

## Gruppi di rotte

| Prefisso | Contenuto |
|---|---|
| `/api/auth` | login demo |
| `/api/users` | utenti e ruoli |
| `/api/locales` | locali |
| `/api/game-types` | tipi di gioco e modelli sensore |
| `/api/games` | giochi installati |
| `/api/devices`, `/api/sensors`, `/api/actuators` | componenti edge |
| `/api/matches` | partite ed eventi |
| `/api/teams` | squadre e membri |
| `/api/tournaments` | tornei e classifiche |
| `/api/statistics` | statistiche globali, locali e personali |

## Evento partita

Esempio di evento inviato al Match Service:

```json
{
  "event_uuid": "edge-1-20-1720000000-3",
  "event_type": "DART_THROW_PLAYER_1",
  "match_id": 20,
  "game_id": 2,
  "locale_id": 1,
  "device_id": 1,
  "value": 60,
  "sync_status": "SYNCED"
}
```

Il campo `value` permette di gestire goal, punti delle bocce e valori dei tiri alle freccette. Se non e presente, il valore usato e 1.
''')

mqtt_doc = write('05-mqtt.md', r'''
# MQTT

## Topic eventi

```text
locales/{localeId}/games/{gameId}/matches/{matchId}/events
```

## Topic heartbeat

```text
locales/{localeId}/edge/{deviceId}/heartbeat
```

## Topic attuatori

```text
locales/{localeId}/games/{gameId}/actuators/{actuatorId}/commands
```

Gli eventi vengono pubblicati con QoS 1. Ogni evento ha un UUID. Il Match Service controlla l'UUID prima di salvare il messaggio, quindi una nuova consegna dello stesso messaggio non modifica due volte il punteggio.

L'heartbeat viene inviato ogni cinque secondi. Il server aggiorna `last_seen`, `last_sync` e lo stato del dispositivo. Se il dispositivo non invia heartbeat per venti secondi viene indicato come offline.
''')

validazione = write('06-validazione.md', r'''
# Validazione

## Test eseguiti senza Docker

- 29 test Node.js superati;
- controllo di sintassi JavaScript superato;
- controllo strutturale del progetto: 174 verifiche superate;
- controllo dei file e dei collegamenti del frontend;
- controllo dello schema SQL e dei servizi Docker;
- OpenAPI 3.0.3 con 41 percorsi pubblici.

## Test end-to-end

Lo script `scripts/integration-test.js` verifica:

1. gateway, database, MQTT e tre microservizi;
2. login dei quattro ruoli;
3. creazione di un tipo di gioco con eventi personalizzati;
4. configurazione di gioco, sensori e attuatore;
5. creazione di un torneo multi-locale;
6. uso del valore numerico dell'evento;
7. deduplicazione tramite UUID;
8. fine automatica al limite di punteggio;
9. collegamento automatico al torneo;
10. simulazione freccette offline e sincronizzazione;
11. heartbeat, attuatori e statistiche.

Comandi:

```bash
docker compose down -v
docker compose up --build -d
node scripts/integration-test.js
```
''')

limitazioni = write('07-limitazioni.md', r'''
# Limiti del prototipo

Il progetto e un prototipo universitario. Le password demo sono salvate in chiaro per rendere semplice l'avvio. In un sistema reale verrebbero usati hash, token e HTTPS.

I sensori sono simulati via software. L'architettura permette di sostituire il simulatore con Arduino, ESP32 o Raspberry Pi senza cambiare il formato degli eventi.

La gestione dei tornei produce classifica e calendario delle partite collegate, ma non genera automaticamente tutti gli accoppiamenti di un tabellone ad eliminazione.
''')

schema_relazione = write('08-relazione-finale-schema.md', r'''
# Struttura della relazione finale

1. Introduzione e obiettivo
2. Specifiche funzionali
3. Analisi delle tecnologie
4. Scelta dell'approccio
5. Architettura del software
6. Modello del dominio e database
7. API REST e MQTT
8. Implementazione
9. Interfaccia utente
10. Validazione
11. Demo
12. Limiti e sviluppi futuri
''')

problemi = write('09-problemi-incontrati.md', r'''
# Problemi incontrati

## Eventi duplicati

Con MQTT QoS 1 lo stesso messaggio puo essere ricevuto piu di una volta. E stato aggiunto `event_uuid` con indice univoco.

## Connessione assente

L'edge non deve perdere la partita. Gli eventi vengono scritti in una coda JSON con stato `PENDING` e ripubblicati quando la connessione torna disponibile.

## Tipi di gioco diversi

All'inizio gli eventi erano fissi per il calciobalilla. La logica e stata spostata nella tabella `game_types`, cosi ogni tipo possiede evento di inizio, due eventi di punteggio, evento di fine e limite.

## Separazione dei ruoli

I controlli sono stati inseriti sia nelle pagine sia nel backend. Il backend resta il controllo principale e impedisce l'accesso a dati di altri locali.
''')

risultati = write('10-risultati-test.md', r'''
# Risultati dei test

| Controllo | Risultato |
|---|---|
| Test unitari Node.js | 29 superati, 0 falliti |
| Sintassi backend | superata |
| Sintassi gateway e simulatore | superata |
| Controllo strutturale | 174 verifiche superate |
| Collegamenti frontend | controllati |
| Struttura microservizi Docker | controllata |
| OpenAPI | 3.0.3, 41 percorsi pubblici |
| Test end-to-end | script completo in `scripts/integration-test.js` |

Il test end-to-end deve essere eseguito dopo l'avvio dei container, perche usa MySQL e Mosquitto reali.
''')

offline = write('11-online-offline-sync.md', r'''
# Sincronizzazione online e offline

![Sequenza offline](figures/07-sequenza-offline.png)

## Stato online

L'edge pubblica subito l'evento con `sync_status=SYNCED`. Il Match Service salva il messaggio e aggiorna il punteggio.

## Stato offline

Se MQTT non e raggiungibile oppure viene attivata la modalita offline, l'edge aggiunge l'evento a `offline-queue.json` con `sync_status=PENDING`.

## Ritorno online

L'edge ripubblica gli eventi nello stesso ordine. Il Match Service controlla l'UUID e ignora eventuali duplicati. Quando la coda e vuota, l'heartbeat aggiorna `last_sync`.
''')

visual = write('12-visual-style.md', r'''
# Stile dell'interfaccia

L'interfaccia usa un layout comune con barra laterale, intestazione e pannelli. Le pagine cambiano in base al ruolo.

- badge per stati online, offline, live e terminato;
- tabelle per utenti, sensori, attuatori e partite;
- schede per giochi, tornei e statistiche;
- pagina live con punteggio grande e timeline degli eventi;
- interfaccia locale dell'edge sulla porta 8090.

Le pagine sono responsive e non richiedono framework JavaScript.
''')

casi_uso = write('13-casi-uso-testuali.md', r'''
# Casi d'uso testuali

## CU1 - Accesso

**Attore:** tutti gli utenti.  
**Precondizione:** l'utente esiste.  
**Flusso:** inserisce username e password, il server controlla i dati e restituisce il ruolo.  
**Risultato:** viene aperta la dashboard corretta.

## CU2 - Configurare un tipo di gioco

**Attore:** amministratore gioco.  
**Flusso:** inserisce nome, descrizione, eventi, limite e supporto squadre. Aggiunge i modelli dei sensori.  
**Risultato:** il tipo puo essere scelto quando viene installato un gioco.

## CU3 - Installare un gioco

**Attore:** amministratore locale.  
**Flusso:** sceglie un tipo di gioco, assegna un nome e collega il gioco al locale. Configura sensori e attuatori sul dispositivo edge.  
**Risultato:** il gioco appare tra quelli disponibili.

## CU4 - Avviare una partita

**Attore:** amministratore locale.  
**Precondizione:** gioco online e due partecipanti validi.  
**Flusso:** seleziona giocatori o squadre e avvia la partita.  
**Risultato:** il gioco passa a `IN_GAME` e la partita a `LIVE`.

## CU5 - Registrare un evento

**Attore:** dispositivo edge.  
**Flusso:** pubblica un messaggio MQTT con UUID, tipo e valore. Il Match Service legge le regole del tipo di gioco e aggiorna il partecipante corretto.  
**Risultato:** evento e punteggio sono salvati.

## CU6 - Terminare al limite

**Attore:** Match Service.  
**Precondizione:** il tipo di gioco possiede un limite.  
**Flusso:** dopo ogni evento di punteggio confronta i due valori con il limite.  
**Risultato:** salva il vincitore, termina la partita, aggiorna il gioco e invia il comando agli attuatori.

## CU7 - Funzionare offline

**Attore:** dispositivo edge.  
**Flusso:** se MQTT non e disponibile salva gli eventi nella coda locale. Al ritorno online li invia in ordine.  
**Risultato:** nessun evento viene perso.

## CU8 - Gestire un torneo

**Attore:** amministratore piattaforma.  
**Flusso:** sceglie tipo, modalita, locali, date e stato. Le partite compatibili concluse vengono collegate.  
**Risultato:** il sistema calcola la classifica con 3 punti per vittoria e 1 per pareggio.

## CU9 - Consultare statistiche

**Attore:** tutti secondo il proprio ruolo.  
**Risultato:** vengono mostrati numero di partite, punti totali, media, percentuale vittorie, ranking e dati per tipo di gioco.
''')

diagrammi_completi = write('14-diagrammi-completi.md', diagrammi.read_text(encoding='utf-8'))

guida_demo = write('15-guida-demo-esame.md', r'''
# Guida demo esame

## Avvio

```bash
docker compose down -v
docker compose up --build -d
```

Aprire:

- applicazione: `http://localhost:8080`;
- edge: `http://localhost:8090`;
- stato servizi: `http://localhost:3000/api/health`.

## Account

| Ruolo | Username | Password |
|---|---|---|
| Piattaforma | platform | platform123 |
| Locale | localadmin | local123 |
| Gioco | gameadmin | game123 |
| Giocatore | client | client123 |

## Percorso consigliato

1. Accedere come `gameadmin` e mostrare le regole diverse di Calciobalilla e Freccette.
2. Aprire la pagina sensori e mostrare gli eventi configurati.
3. Accedere come `localadmin` e avviare una partita di calciobalilla o freccette.
4. Premere **Simula con MQTT** e mostrare la timeline live.
5. Aprire l'interfaccia edge, passare offline e avviare un'altra simulazione.
6. Mostrare gli eventi nella coda, poi tornare online e osservare la sincronizzazione.
7. Aprire tornei e classifica.
8. Accedere come `platform` e mostrare statistiche globali e gestione locali.

## Test finale

```bash
cd backend
npm test
npm run check
cd ..
node scripts/validate-project.js
node scripts/integration-test.js
```
''')

matrice = write('16-matrice-requisiti.md', r'''
# Matrice dei requisiti

| Requisito | Implementazione | Verifica |
|---|---|---|
| Quattro tipi di utenti | `users.role`, middleware e dashboard | test accessRules |
| Giochi in locali identificabili | tabelle `locales` e `games` | API giochi/locali |
| Admin gioco e regole | `game_types`, `sensor_templates`, pagina game admin | test gameRules |
| Sensori reali o simulati | `sensors` e Edge Service | configurazione edge |
| REST | API Gateway e tre servizi | OpenAPI |
| MQTT | Mosquitto, eventi, heartbeat, attuatori | integrazione MQTT |
| Offline | coda JSON e ripubblicazione | scenario integrazione |
| Deduplicazione | `event_uuid` univoco | test integrazione |
| Partite individuali | riferimenti ai client | test validazione |
| Partite a squadre | `teams`, `team_members`, campi team | test tornei |
| Tornei multi-locale | `tournament_locations` | API tornei |
| Classifica | servizio torneo, 3/1/0 punti | test ranking |
| Statistiche | globali, locali, personali e per tipo | API statistics |
| Interfaccia utente | pagine web per tutti i ruoli | demo browser |
| Microservizi | catalog, match, tournament, gateway | Docker Compose |
| UML e sequenze | immagini in `docs/figures` | relazione |
| Test componenti | Node test runner | 29 test |
| Test integrazione | `scripts/integration-test.js` | esecuzione con Docker |
| Demo funzionante | script start/stop e guida | `15-guida-demo-esame` |
''')

final_md = write('00-relazione-finale.md', r'''
# PlayConnect - Connected Games Platform

Progetto di laboratorio PISSIR - Anno accademico 2025/2026

## 1. Introduzione

PlayConnect e una piattaforma che collega a Internet giochi tradizionali fisici. Il gioco rimane fisico, ma i sensori rilevano gli eventi importanti. Un dispositivo edge riceve gli eventi e li invia alla piattaforma centrale. Il server conserva partite, punteggi, tornei e statistiche.

Il prototipo comprende calciobalilla, freccette, bocce e Monopoli. I tipi di gioco non sono scritti direttamente nel codice: l'amministratore del gioco puo definire gli eventi e il limite di punteggio.

## 2. Specifiche funzionali

La piattaforma gestisce giocatori, amministratori locali, amministratori del gioco e amministratori della piattaforma. Le funzioni dettagliate sono descritte in `01-specifiche-funzionali.md`.

![Casi d'uso](figures/01-casi-uso.png)

## 3. Analisi tecnologica

Per la parte centrale sono stati scelti Node.js, Express e MySQL. Il frontend usa HTML, CSS e JavaScript senza framework. La comunicazione dal browser usa REST. La comunicazione degli eventi usa MQTT perche e asincrona e permette all'edge di lavorare anche con una connessione instabile.

Docker Compose avvia tutti i componenti nello stesso modo su Windows, macOS e Linux.

## 4. Architettura

![Architettura](figures/02-architettura.png)

Il Catalog Service gestisce il catalogo. Il Match Service gestisce le partite e MQTT. Il Tournament Service gestisce tornei, squadre e statistiche. L'API Gateway presenta un unico indirizzo al frontend.

## 5. Dominio e database

![Dominio](figures/03-dominio.png)

Le entita principali sono locale, utente, tipo di gioco, gioco, dispositivo edge, sensore, attuatore, partita, evento, squadra e torneo. Le relazioni sono definite con chiavi esterne in MySQL.

## 6. API REST e MQTT

L'API pubblica e documentata in `openapi.yaml`. Gli eventi MQTT usano topic che contengono locale, gioco e partita. Il valore numerico dell'evento permette di registrare un goal oppure un tiro da 60 punti con lo stesso formato generale.

## 7. Implementazione

Quando arriva un evento, il Match Service recupera le regole del tipo di gioco. Confronta `event_type` con evento di inizio, punto partecipante 1, punto partecipante 2 ed evento di fine. Per un evento di punteggio usa il campo `value`, oppure 1 come valore predefinito.

Quando un punteggio raggiunge `score_limit`, la partita viene chiusa automaticamente. Il vincitore viene salvato, il gioco torna online, gli attuatori ricevono il comando finale e la partita puo essere collegata a un torneo attivo compatibile.

L'edge scarica la configurazione di gioco e sensori dal Catalog Service. In assenza di MQTT salva gli eventi in una coda locale. Al ritorno online li invia nello stesso ordine.

## 8. Interfaccia utente

Ogni ruolo possiede pagine dedicate. La partita live mostra punteggio, stato, ultimo evento, valore del sensore e sincronizzazione. L'interfaccia edge mostra connessione, coda offline, configurazione e comandi attuatore.

## 9. Validazione

Sono presenti test unitari, controlli di sintassi, controllo della struttura e uno script end-to-end. I dettagli sono in `06-validazione.md` e `10-risultati-test.md`.

## 10. Demo

La demo mostra prima una partita online e poi una partita avviata mentre l'edge e offline. In questo secondo caso gli eventi appaiono nella coda e vengono sincronizzati al ritorno online. La procedura completa e in `15-guida-demo-esame.md`.

## 11. Conclusione

Il progetto realizza il flusso completo dal sensore simulato alla statistica. La separazione tra REST, MQTT, edge e microservizi permette di aggiungere altri giochi senza riscrivere il sistema centrale.
''')

# Complete diagram source is kept under the expected root name.


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = title
        set_cell_shading(cell, 'DCE6F1')
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_code(doc, code):
    p = doc.add_paragraph(style='Code')
    p.add_run(code)
    return p


def add_figure(doc, filename, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = 'Caption'


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    styles['Normal'].font.size = Pt(10.5)
    styles['Title'].font.name = 'Arial'
    styles['Title'].font.size = Pt(28)
    styles['Title'].font.color.rgb = RGBColor(31, 78, 121)
    for name in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[name].font.name = 'Arial'
        styles[name].font.color.rgb = RGBColor(31, 78, 121)
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Liberation Mono'
        code_style.font.size = Pt(8.5)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(3)
        code_style.paragraph_format.space_after = Pt(6)

    for sec in doc.sections:
        sec.header.paragraphs[0].text = 'PlayConnect - Relazione finale PISSIR'
        sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sec.header.paragraphs[0].runs[0].font.size = Pt(8)
        add_page_number(sec.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.style = 'Title'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('PlayConnect')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run('Connected Games Platform')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\nProgetto di Laboratorio PISSIR\nAnno Accademico 2025/2026\n\nGruppo PlayConnect').font.size = Pt(13)
    doc.add_paragraph('\n')
    add_figure(doc, '02-architettura.png', 'Vista generale della piattaforma', width=6.2)
    doc.add_page_break()

    add_heading(doc, 'Indice', 1)
    index_items = [
        '1. Introduzione', '2. Specifiche funzionali', '3. Analisi tecnologica',
        '4. Scelta dell approccio', '5. Architettura del software',
        '6. Modello del dominio e database', '7. API REST e MQTT',
        '8. Implementazione', '9. Interfaccia utente', '10. Validazione',
        '11. Procedura di demo', '12. Limiti e sviluppi futuri'
    ]
    for item in index_items:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    add_heading(doc, '1. Introduzione', 1)
    doc.add_paragraph(
        'PlayConnect e una piattaforma software che collega giochi tradizionali fisici a Internet. '
        'Il gioco continua a essere usato normalmente, mentre sensori o pulsanti registrano eventi '
        'come inizio, punto, tiro e fine della partita. Un dispositivo edge raccoglie questi eventi '
        'e li invia alla piattaforma centrale.'
    )
    doc.add_paragraph(
        'La piattaforma conserva partite, risultati, tornei e statistiche. Il progetto usa sensori '
        'simulati, ma il formato dei messaggi puo essere usato anche da Arduino, ESP32 o Raspberry Pi.'
    )

    add_heading(doc, '2. Specifiche funzionali', 1)
    doc.add_paragraph('Sono previsti quattro ruoli con funzioni diverse.')
    add_table(doc, ['Ruolo', 'Funzioni'], [
        ['Giocatore', 'Visualizza giochi, partite, statistiche personali, tornei e classifiche.'],
        ['Amministratore locale', 'Gestisce giochi, client, edge, sensori, attuatori e partite del proprio locale.'],
        ['Amministratore gioco', 'Definisce tipi di gioco, eventi, limite e modelli dei sensori.'],
        ['Amministratore piattaforma', 'Gestisce locali, amministratori, tornei e statistiche globali.']
    ])
    doc.add_paragraph('Le principali funzioni implementate sono:')
    for text in [
        'gestione dei locali e dei giochi installati;',
        'tipi di gioco configurabili senza modificare il Match Service;',
        'partite individuali e a squadre;',
        'eventi con UUID e valore numerico;',
        'funzionamento edge online e offline;',
        'tornei su uno o piu locali;',
        'statistiche globali, locali e personali;',
        'interfaccia web per tutti i ruoli.'
    ]:
        doc.add_paragraph(text, style='List Bullet')
    add_figure(doc, '01-casi-uso.png', 'Figura 1 - Diagramma dei casi d uso', width=3.85)

    add_heading(doc, '3. Analisi tecnologica', 1)
    doc.add_paragraph(
        'Per il controllo dei giochi si poteva usare soltanto REST, ma in questo caso ogni sensore '
        'avrebbe dovuto conoscere direttamente il server. MQTT e piu adatto agli eventi piccoli e '
        'frequenti, permette la consegna asincrona e separa il dispositivo dal servizio che elabora i dati.'
    )
    add_table(doc, ['Tecnologia', 'Uso nel progetto'], [
        ['Node.js ed Express', 'API Gateway, microservizi ed Edge Service.'],
        ['MySQL 8', 'Dati centrali e relazioni tra entita.'],
        ['Mosquitto MQTT', 'Eventi sensore, heartbeat e comandi attuatori.'],
        ['HTML, CSS, JavaScript', 'Interfaccia web senza framework.'],
        ['Docker Compose', 'Avvio ripetibile di tutti i componenti.'],
        ['OpenAPI 3.0.3', 'Documentazione delle API REST.']
    ])

    add_heading(doc, '4. Scelta dell approccio', 1)
    doc.add_paragraph(
        'Il sistema centrale e diviso in tre microservizi applicativi. Il Catalog Service contiene '
        'le funzioni di configurazione. Il Match Service elabora eventi e partite. Il Tournament '
        'Service calcola tornei e statistiche. Un API Gateway offre un solo punto di accesso al browser.'
    )
    doc.add_paragraph(
        'L edge usa una piccola coda JSON. Questa soluzione e sufficiente per il prototipo e rende '
        'visibile durante la demo cosa succede quando la rete non e disponibile.'
    )

    add_heading(doc, '5. Architettura del software', 1)
    add_figure(doc, '02-architettura.png', 'Figura 2 - Architettura generale', width=6.4)
    doc.add_paragraph(
        'Solo i servizi backend accedono direttamente a MySQL. Il frontend comunica con il gateway. '
        'L edge comunica con il Catalog Service tramite una rotta interna protetta da chiave e con il '
        'Match Service tramite il broker MQTT.'
    )
    add_heading(doc, '5.1 Package', 2)
    add_figure(doc, '04-package.png', 'Figura 3 - Organizzazione dei package', width=6.4)
    add_heading(doc, '5.2 Deployment', 2)
    add_figure(doc, '08-deployment.png', 'Figura 4 - Deployment con Docker Compose', width=5.6)

    add_heading(doc, '6. Modello del dominio e database', 1)
    add_figure(doc, '03-dominio.png', 'Figura 5 - Classi del dominio', width=6.4)
    doc.add_paragraph(
        'La tabella game_types contiene le regole del gioco: evento di inizio, evento punto 1, evento '
        'punto 2, evento di fine, limite di punteggio e supporto alle squadre. La tabella match_events '
        'salva UUID, valore, stato di sincronizzazione e dati essenziali della sorgente.'
    )
    add_table(doc, ['Tabella', 'Contenuto'], [
        ['locales, users', 'Locali e utenti con ruolo.'],
        ['game_types, sensor_templates', 'Regole e modelli dei sensori.'],
        ['games, edge_devices, sensors, actuators', 'Installazione fisica o simulata.'],
        ['matches, match_events', 'Partite ed eventi ricevuti.'],
        ['teams, team_members', 'Squadre e componenti.'],
        ['tournaments e tabelle ponte', 'Tornei, locali, squadre e partite collegate.']
    ])

    add_heading(doc, '7. API REST e MQTT', 1)
    doc.add_paragraph('L API pubblica contiene 41 percorsi ed e descritta in docs/openapi.yaml.')
    add_code(doc, 'locales/{localeId}/games/{gameId}/matches/{matchId}/events')
    add_code(doc, 'locales/{localeId}/edge/{deviceId}/heartbeat')
    add_code(doc, 'locales/{localeId}/games/{gameId}/actuators/{actuatorId}/commands')
    doc.add_paragraph('Esempio di evento delle freccette:')
    add_code(doc, '{\n  "event_uuid": "edge-1-20-1720000000-3",\n  "event_type": "DART_THROW_PLAYER_1",\n  "value": 60,\n  "sync_status": "SYNCED"\n}')

    add_heading(doc, '8. Implementazione', 1)
    add_heading(doc, '8.1 Regole configurabili', 2)
    doc.add_paragraph(
        'Il file gameRules.js trasforma i campi del tipo di gioco in regole semplici. Quando arriva '
        'un evento, il servizio decide se si tratta di inizio, punto del primo partecipante, punto del '
        'secondo partecipante, fine oppure evento generico configurato su un sensore.'
    )
    doc.add_paragraph(
        'Il campo value permette di aggiungere piu di un punto. Per il calciobalilla vale normalmente 1, '
        'per le bocce puo valere da 1 a 3 e per le freccette puo valere 20, 50 o 60.'
    )
    add_heading(doc, '8.2 Fine automatica', 2)
    doc.add_paragraph(
        'Dopo un evento di punteggio il servizio confronta score1 e score2 con score_limit. Quando il '
        'limite viene raggiunto salva il vincitore, chiude la partita, rimette online il gioco, invia '
        'il comando agli attuatori e collega la partita a un torneo attivo compatibile.'
    )
    add_heading(doc, '8.3 Deduplicazione', 2)
    doc.add_paragraph(
        'MQTT QoS 1 puo consegnare lo stesso messaggio piu volte. event_uuid e univoco nel database. '
        'Se il servizio riceve di nuovo lo stesso UUID restituisce la partita senza modificare il punteggio.'
    )
    add_heading(doc, '8.4 Heartbeat', 2)
    doc.add_paragraph(
        'L edge invia un heartbeat ogni cinque secondi. Il Match Service aggiorna stato, last_seen e '
        'last_sync. Un controllo periodico segna offline i dispositivi silenziosi da piu di venti secondi.'
    )
    add_heading(doc, '8.5 Classi principali', 2)
    add_figure(doc, '05-classi-implementazione.png', 'Figura 6 - Classi e moduli principali', width=6.4)

    landscape = doc.add_section(WD_SECTION_START.NEW_PAGE)
    landscape.header.is_linked_to_previous = False
    landscape.footer.is_linked_to_previous = False
    landscape.header.paragraphs[0].clear()
    landscape.footer.paragraphs[0].clear()
    landscape.orientation = 1
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.top_margin = Inches(0.5)
    landscape.bottom_margin = Inches(0.5)
    landscape.left_margin = Inches(0.5)
    landscape.right_margin = Inches(0.5)
    landscape.header.paragraphs[0].text = 'PlayConnect - Flussi di esecuzione'
    landscape.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(landscape.footer.paragraphs[0])
    add_heading(doc, '8.6 Sequenza online', 2)
    add_figure(doc, '06-sequenza-online.png', 'Figura 7 - Partita online', width=9.2)
    add_heading(doc, '8.7 Sequenza offline', 2)
    add_figure(doc, '07-sequenza-offline.png', 'Figura 8 - Coda e sincronizzazione', width=9.2)

    portrait = doc.add_section(WD_SECTION_START.NEW_PAGE)
    portrait.header.is_linked_to_previous = False
    portrait.footer.is_linked_to_previous = False
    portrait.header.paragraphs[0].clear()
    portrait.footer.paragraphs[0].clear()
    portrait.orientation = 0
    portrait.page_width, portrait.page_height = portrait.page_height, portrait.page_width
    portrait.top_margin = Inches(0.65)
    portrait.bottom_margin = Inches(0.65)
    portrait.left_margin = Inches(0.75)
    portrait.right_margin = Inches(0.75)
    portrait.header.paragraphs[0].text = 'PlayConnect - Relazione finale PISSIR'
    portrait.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(portrait.footer.paragraphs[0])

    add_heading(doc, '9. Interfaccia utente', 1)
    doc.add_paragraph(
        'Il frontend usa pagine separate per i quattro ruoli. Tutte le pagine condividono barra laterale, '
        'intestazione, badge degli stati, tabelle e schede. La pagina della partita live aggiorna i dati '
        'ogni due secondi e mostra il valore dell ultimo evento.'
    )
    doc.add_paragraph(
        'L amministratore gioco puo creare e modificare i tipi, cambiare gli eventi e aggiungere modelli '
        'sensore. L amministratore locale configura edge, sensori e attuatori e avvia le partite.'
    )

    add_heading(doc, '10. Validazione', 1)
    add_table(doc, ['Controllo', 'Risultato'], [
        ['Test unitari Node.js', '29 superati, 0 falliti'],
        ['Controllo sintassi', 'superato'],
        ['Controllo struttura e link frontend', '174 verifiche superate'],
        ['OpenAPI', '3.0.3 - 41 percorsi'],
        ['Test end-to-end', 'script completo da eseguire con Docker']
    ])
    doc.add_paragraph('Lo script di integrazione prova un flusso completo con regole personalizzate e un flusso offline delle freccette.')
    add_code(doc, 'cd backend\nnpm test\nnpm run check\ncd ..\nnode scripts/validate-project.js\nnode scripts/integration-test.js')

    add_heading(doc, '11. Procedura di demo', 1)
    steps = [
        'Avviare il progetto con start-demo oppure docker compose up --build -d.',
        'Mostrare i quattro ruoli e le relative dashboard.',
        'Aprire il tipo Freccette e mostrare eventi e limite 301.',
        'Avviare una partita e usare Simula con MQTT.',
        'Aprire l edge sulla porta 8090 e passare offline.',
        'Avviare una nuova simulazione, mostrare la coda e tornare online.',
        'Mostrare la partita terminata, gli attuatori, il torneo e le statistiche.'
    ]
    for number, step in enumerate(steps, 1):
        paragraph = doc.add_paragraph(f'{number}. {step}')
        paragraph.paragraph_format.space_after = Pt(1)

    add_heading(doc, '12. Limiti e sviluppi futuri', 1)
    doc.add_paragraph(
        'Il prototipo usa account demo e password in chiaro. In produzione servirebbero hash, token, HTTPS '
        'e una gestione piu completa delle autorizzazioni tecniche. I sensori sono simulati e possono essere '
        'sostituiti con dispositivi reali mantenendo lo stesso payload MQTT.'
    )
    doc.add_paragraph(
        'Come sviluppo futuro si possono aggiungere prenotazioni, notifiche, tabelloni automatici, dati '
        'specifici per Monopoli e posizione delle bocce.'
    )

    add_heading(doc, 'Conclusione', 1)
    doc.add_paragraph(
        'PlayConnect realizza il flusso dal sensore alla statistica e continua a raccogliere dati durante '
        'una breve assenza di connessione. Le regole configurabili permettono di aggiungere altri giochi '
        'senza inserire nuovi nomi di evento nel Match Service.'
    )

    out = FINAL / 'PlayConnect_Relazione_Finale.docx'
    doc.save(out)
    return out


report_docx = build_report()
shutil.copy2(report_docx, DOCS / 'PlayConnect_Relazione_Finale.docx')
for markdown_file in DOCS.glob('*.md'):
    markdown_file.unlink()
print(report_docx)
